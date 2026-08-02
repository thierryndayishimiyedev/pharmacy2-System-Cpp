from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_BREAK
from pathlib import Path

OUT=Path('Final_Deliverables'); OUT.mkdir(exist_ok=True)
doc=Document(); sec=doc.sections[0]
for s in doc.sections:
    s.top_margin=s.bottom_margin=s.left_margin=s.right_margin=Inches(1)
    s.header_distance=s.footer_distance=Inches(.45)
styles=doc.styles
normal=styles['Normal']; normal.font.name='Times New Roman'; normal._element.rPr.rFonts.set(qn('w:eastAsia'),'Times New Roman'); normal.font.size=Pt(12); normal.paragraph_format.line_spacing=1.5; normal.paragraph_format.space_after=Pt(6)
for nm,size,color in [('Title',22,'17365D'),('Heading 1',15,'17365D'),('Heading 2',13,'1F4E79'),('Heading 3',12,'1F4E79')]:
    st=styles[nm]; st.font.name='Times New Roman'; st._element.rPr.rFonts.set(qn('w:eastAsia'),'Times New Roman'); st.font.size=Pt(size); st.font.bold=True; st.font.color.rgb=RGBColor.from_string(color); st.paragraph_format.space_before=Pt(12); st.paragraph_format.space_after=Pt(6)
mono=styles.add_style('Diagram', WD_STYLE_TYPE.PARAGRAPH); mono.font.name='Courier New'; mono.font.size=Pt(9); mono.paragraph_format.line_spacing=1.0
def page_num(p):
    r=p.add_run(); fld=OxmlElement('w:fldSimple'); fld.set(qn('w:instr'),'PAGE'); r._r.addnext(fld)
def footer(section):
    p=section.footer.paragraphs[0]; p.alignment=WD_ALIGN_PARAGRAPH.RIGHT; p.add_run('Motherland Pharmacy Management System | Page '); page_num(p)
def p(t='',bold=False,align=None,style=None):
    q=doc.add_paragraph(style=style); q.alignment=align or WD_ALIGN_PARAGRAPH.LEFT; r=q.add_run(t); r.bold=bold; return q
def h(t,l=1): doc.add_heading(t,level=l)
def bullet(t): doc.add_paragraph(t,style='List Bullet')
def table(headers,rows,widths=None):
    t=doc.add_table(rows=1, cols=len(headers)); t.style='Table Grid'; t.alignment=WD_TABLE_ALIGNMENT.CENTER
    for i,x in enumerate(headers):
        c=t.rows[0].cells[i]; c.text=x; c.vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.CENTER
        for r in c.paragraphs[0].runs: r.bold=True
    for row in rows:
        cells=t.add_row().cells
        for i,x in enumerate(row): cells[i].text=str(x); cells[i].vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.CENTER
    return t
def caption(txt): p('Figure: '+txt,align=WD_ALIGN_PARAGRAPH.CENTER)

# cover
for text,size,bold in [('RP KARONGI COLLEGE',16,True),('DEPARTMENT OF ICT',14,True),('PROGRAM: INFORMATION TECHNOLOGY (IT)',13,True),('',12,False),('CAT PRACTICE: PROJECT BASED',14,True),('',12,False),('MOTHERLAND PHARMACY MANAGEMENT SYSTEM',20,True),('',12,False),('COURSE: DATA STRUCTURES AND ALGORITHMS',13,True),('CASE STUDY: MOTHERLAND PHARMACY',13,True),('',12,False),('PREPARED BY:',12,True),('THIERRY NDAYISHIMIYE',14,True),('',12,False),('INSTITUTION: RP KARONGI COLLEGE',12,True),('DATE: 1 August 2026',12,True)]:
    p(text,bold,WD_ALIGN_PARAGRAPH.CENTER)
doc.add_page_break(); footer(doc.sections[0])

h('Table of Contents');
for x in ['1. Description of the Project','2. Problem Statement','3. Objectives','4. System Requirements','5. Algorithm Design','6. Database Design','7. System Implementation','8. Conclusion and Recommendations','References','Appendix A: Screenshot Placeholders']: p(x)
doc.add_page_break()
h('1. Description of the Project')
p('Motherland Pharmacy Management System is a terminal-based C11 application connected directly to the local MySQL/MariaDB database named pharmacy_db. It was implemented for the Motherland Pharmacy case study to centralize day-to-day pharmacy records. The program is menu driven and displays the currently signed-in user in its Motherland Pharmacy header.')
p('The application connects through MySQL Connector/C. database.c initializes a MYSQL handle and connects to 127.0.0.1 on port 3306 using the configured root account and pharmacy_db. The program stores and retrieves records with SQL queries and presents results in the console.')
p('The implemented users are pharmacy staff with accounts in the users table. Roles stored by the system are Admin and Pharmacist; however, the source code does not restrict menu access by role. The main modules are authentication, user management, supplier management, customer management, medicine inventory, sales, reports, dashboard, and input/output validation.')
h('2. Problem Statement')
p('Pharmacy operations require reliable records of medicines, stock quantity, expiry dates, suppliers, customers, and completed sales. When these records are managed manually or retrieved slowly, staff can face inconsistent quantities, difficulty finding records, delayed reports, and a risk of selling unsuitable stock. Motherland Pharmacy requires a local system that records these data items consistently, reduces stock after a sale, supports safe expiry handling, and produces operational information from stored transactions.')
h('3. Objectives')
h('3.1 General Objective',2); p('To develop a C and MySQL based pharmacy management system that supports reliable medicine, sales, customer, supplier, and user record management for Motherland Pharmacy.')
h('3.2 Specific Objectives',2)
for x in ['Authenticate staff before dashboard access.','Maintain users, suppliers, customers, and medicines using add, view, search, update, and delete operations.','Record multi-item medicine sales and update stock transactionally.','Prevent sale of expired medicines and enforce basic price, quantity, date, and contact validation.','Generate dashboard, sales, inventory, expiry, and advisory reports.']: bullet(x)
h('4. System Requirements')
h('4.1 Functional Requirements',2)
table(['Requirement','Actual behavior'],[
('Authentication','Validates username/password against users and stores currentUserId/display name.'),('Master data','CRUD menus for users, suppliers, customers, and medicines.'),('Sales','Supports walk-in or registered customer, a multi-item cart, receipt view, search, and sale deletion.'),('Stock safety','Checks available quantity; blocks expired medicine; commits/rolls back sale database work.'),('Reports','Provides daily/all-time/date/monthly sales, low stock, expiry, top-selling, stock value, dashboard, and advice.'),('Cleanup','Deletes only expired medicines without sale-item history after DELETE confirmation.')])
h('4.2 Non-Functional Requirements',2)
table(['Area','Implementation-oriented requirement'],[('Performance','Local SQL queries should return practical records promptly for a small local deployment.'),('Reliability','Sales and deletion use database transactions.'),('Usability','Console menus, formatted status messages, and short ID lists guide staff.'),('Security','Login and validation exist; plain-text passwords and lack of role authorization are limitations.'),('Maintainability','Modules are split into .c/.h files with shared UI/database interfaces.'),('Availability','Requires a running local XAMPP MySQL/MariaDB service.'),('Integrity','Foreign keys and application validation protect key relationships.')])
h('5. Algorithm Design')
h('5.1 Main Pseudocode',2)
p('START\nCONNECT to pharmacy_db\nIF connection fails THEN display error and EXIT\nREPEAT show Main Menu; read choice\n  IF Login THEN validate credentials; if valid show Dashboard\n  Dashboard routes to CRUD, Sales, Reports, or System Overview\n  IF Logout return to Main Menu\n  IF Exit disconnect database and terminate\nEND REPEAT',style='Diagram')
h('5.2 Sales Pseudocode',2)
p('Select customer or walk-in; verify non-walk-in customer. Repeat: select medicine; query price, stock and expiry; reject expired/invalid/insufficient item; add or merge cart. Confirm cart. START TRANSACTION. Insert sales header. For each cart item insert sale_items and decrement medicines quantity. On error ROLLBACK; otherwise COMMIT.',style='Diagram')
h('5.3 Main Workflow Flowchart',2)
p('Start -> connectDB -> [failure] print error/exit\n                 -> [success] Main Menu -> Login -> [invalid] Main Menu\n                                              -> [valid] Dashboard -> modules -> Dashboard\n                                                                -> Logout -> Main Menu\nMain Menu -> Exit -> disconnectDB -> End',style='Diagram'); caption('Main system workflow derived from main.c, login.c, and module menus.')
h('6. Database Design')
p('The supplied pharmacy_db.sql defines six InnoDB tables: users, customers, suppliers, medicines, sales, and sale_items. The application uses these tables directly through MySQL Connector/C.')
table(['Table','Purpose','Primary key'],[('users','Staff credentials and role','user_id'),('customers','Customer contact records','customer_id'),('suppliers','Supplier contact records','supplier_id'),('medicines','Inventory, prices, supplier and expiry','medicine_id'),('sales','Sale header, customer/user/date/total','sale_id'),('sale_items','Medicine lines belonging to a sale','sale_item_id')])
h('6.1 ER Diagram',2)
p('SUPPLIERS (supplier_id PK) 1 --- 0..* MEDICINES (supplier_id FK)\nCUSTOMERS (customer_id PK) 1 --- 0..* SALES (customer_id FK, nullable)\nUSERS (user_id PK) 1 --- 0..* SALES (user_id FK, nullable)\nSALES (sale_id PK) 1 --- 1..* SALE_ITEMS (sale_id FK)\nMEDICINES (medicine_id PK) 1 --- 0..* SALE_ITEMS (medicine_id FK)',style='Diagram'); caption('ERD based only on pharmacy_db.sql. Supplier/customer/user references can be set NULL by database delete rules; sale-item parent delete rules are CASCADE.')
h('6.2 Important Columns and Relationships',2)
table(['Entity','Key attributes / relationship'],[('users','username is UNIQUE; role enum Admin/Pharmacist; sales.user_id references users.'),('customers','sales.customer_id references customers and may be NULL for walk-in sales.'),('suppliers','medicines.supplier_id references suppliers; on delete it becomes NULL.'),('medicines','buy_price, sell_price, quantity, expiry_date; sale_items.medicine_id references medicines.'),('sales','grand_total and sale_date; contains one or more sale items in recorded sale workflow.'),('sale_items','quantity, price, subtotal; child of sales and references medicine.')])
h('7. System Implementation')
p('The system is implemented in C and compiled as C11. It has been developed/tested in VS Code and configured for Embarcadero Dev-C++ 6.3. The active connector is MySQL Connector/C 6.1, using mysql.h, libmysql.lib, and libmysql.dll. The database is managed through XAMPP MySQL/MariaDB.')
h('7.1 Module Summary',2)
table(['Module','Implemented role'],[('main/database/login/ui','Startup, connection, login/session, branded output and validators.'),('users/suppliers/customers/medicine','CRUD operations and aggregate list output.'),('sale','Cart, expired/stock checks, transaction, receipts, deletion/restock.'),('report/dashboard_report','Reports, smart advice, safe cleanup and high-level overview.')])
h('7.2 Screenshot Placeholders',2)
p('The workspace did not provide a live running console capture mechanism. Capture the actual screens listed in Appendix A after starting XAMPP; replace each placeholder with the real console screenshot and retain its caption.')
for x in ['Login screen','Motherland Pharmacy dashboard with signed-in user','Manage Medicines and price validation','Sell Medicine cart / expired medicine rejection','Sale receipt','Reports and Smart Pharmacy Advice','System Overview dashboard','Safe expired auto-cleanup confirmation']:
    p('[SCREENSHOT PLACEHOLDER: '+x+']',align=WD_ALIGN_PARAGRAPH.CENTER); caption(x+'.')
h('8. Conclusion and Recommendations')
h('8.1 Conclusion',2); p('The project achieves a local, database-backed pharmacy workflow using C programming, data structures in the form of an in-memory cart array, conditional logic, loops, functions, SQL relations, and transaction control. It provides records and reporting that address medicine, stock, sales, customer, supplier, and staff information requirements for the stated case study.')
h('8.2 Current Features versus Future Recommendations',2)
table(['Current features','Future recommendations'],[('Console CRUD, login, cart sales, reports, expired-sale block, validation, dashboard.','Graphical UI, role-based authorization, password hashing, prepared statements, backups, barcode scanning, batch/purchase records, notification service, cloud database, export reports.')])
h('References')
for x in ['Project C source files: main.c, database.c, login.c, ui.c, users.c, supplier.c, customer.c, medicine.c, sale.c, report.c, dashboard_report.c.','pharmacy_db.sql database schema and supplied sample data.','MySQL Connector/C 6.1 files configured by the project.','DEV_CPP_SETUP.md and motherland_pharmacy.dev in the workspace.']: p(x)
h('Appendix A: Manual Screenshot Checklist')
p('Capture login, dashboard, user/supplier/customer/medicine management, a valid sale cart, expired-sale rejection, receipt, dashboard, report/advice, and safe cleanup screens. Do not show real passwords or sensitive personal contacts in the final report.')
doc.save(OUT/'Motherland_Pharmacy_Management_System_CAT_Report.docx')
